import PyPDF2
from PIL import Image
from docx import Document
import openpyxl
import logging
import mimetypes
from pathlib import Path
from typing import Dict, Any, Optional, List
import json
import io
import base64

logger = logging.getLogger(__name__)

class FileProcessorError(Exception):
    """Custom exception for file processing errors"""
    pass

class FileProcessor:
    """
    Main file processor class that handles multiple file types
    """
    
    def __init__(self):
        self.supported_types = {
            'application/pdf': self.process_pdf,
            'image/jpeg': self.process_image,
            'image/jpg': self.process_image,
            'image/png': self.process_image,
            'image/gif': self.process_image,
            'image/bmp': self.process_image,
            'image/webp': self.process_image,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self.process_docx,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self.process_xlsx,
            'text/plain': self.process_text,
            'text/csv': self.process_csv,
            'application/json': self.process_json
        }
    
    def get_file_type(self, file_path: Path) -> str:
        """
        Determine file type from path and content
        """
        # Try to guess from filename first
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        if mime_type:
            return mime_type
        
        # Fallback to extension-based detection
        suffix = file_path.suffix.lower()
        extension_map = {
            '.pdf': 'application/pdf',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg', 
            '.png': 'image/png',
            '.gif': 'image/gif',
            '.bmp': 'image/bmp',
            '.webp': 'image/webp',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.txt': 'text/plain',
            '.csv': 'text/csv',
            '.json': 'application/json'
        }
        
        return extension_map.get(suffix, 'application/octet-stream')
    
    def is_supported(self, file_path: Path) -> bool:
        """
        Check if file type is supported for processing
        """
        file_type = self.get_file_type(file_path)
        return file_type in self.supported_types
    
    def process_file(self, file_path: Path) -> Dict[str, Any]:
        """
        Process a file and extract content based on its type
        """
        try:
            file_type = self.get_file_type(file_path)
            
            if not self.is_supported(file_path):
                raise FileProcessorError(f"Unsupported file type: {file_type}")
            
            # Get file stats
            stat = file_path.stat()
            
            # Process the file
            processor = self.supported_types[file_type]
            content_data = processor(file_path)
            
            # Return standardized result
            result = {
                'file_path': str(file_path),
                'file_name': file_path.name,
                'file_type': file_type,
                'file_size': stat.st_size,
                'processed_at': stat.st_mtime,
                'processing_status': 'success',
                'content': content_data
            }
            
            logger.info(f"Successfully processed {file_path.name} ({file_type})")
            return result
            
        except Exception as e:
            error_msg = f"Error processing {file_path.name}: {str(e)}"
            logger.error(error_msg)
            
            return {
                'file_path': str(file_path),
                'file_name': file_path.name,
                'file_type': file_type if 'file_type' in locals() else 'unknown',
                'file_size': stat.st_size if 'stat' in locals() else 0,
                'processing_status': 'error',
                'error': error_msg,
                'content': None
            }
    
    def process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text and metadata from PDF files
        """
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                metadata = {}
                if reader.metadata:
                    for key, value in reader.metadata.items():
                        metadata[key] = str(value)
                
                # Extract text from all pages
                text_content = ""
                pages_info = []
                
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    text_content += page_text + "\n\n"
                    
                    pages_info.append({
                        'page_number': page_num + 1,
                        'text_length': len(page_text),
                        'text_preview': page_text[:200] + "..." if len(page_text) > 200 else page_text
                    })
                
                return {
                    'type': 'pdf',
                    'text': text_content.strip(),
                    'page_count': len(reader.pages),
                    'pages': pages_info,
                    'metadata': metadata,
                    'text_length': len(text_content)
                }
                
        except Exception as e:
            raise FileProcessorError(f"PDF processing failed: {str(e)}")
    
    def process_image(self, file_path: Path) -> Dict[str, Any]:
        """
        Process images and extract metadata
        """
        try:
            with Image.open(file_path) as img:
                # Get image info
                image_info = {
                    'type': 'image',
                    'width': img.width,
                    'height': img.height,
                    'mode': img.mode,
                    'format': img.format,
                    'has_transparency': img.mode in ('RGBA', 'LA') or 'transparency' in img.info
                }
                
                # Extract EXIF data if available
                exif_data = {}
                if hasattr(img, '_getexif') and img._getexif():
                    exif_data = {str(k): str(v) for k, v in img._getexif().items()}
                
                # Convert image to base64 for small images (under 1MB)
                image_data = None
                if file_path.stat().st_size < 1024 * 1024:  # 1MB
                    img_byte_arr = io.BytesIO()
                    img.save(img_byte_arr, format=img.format or 'PNG')
                    img_byte_arr = img_byte_arr.getvalue()
                    image_data = base64.b64encode(img_byte_arr).decode('utf-8')
                
                return {
                    **image_info,
                    'exif': exif_data,
                    'image_data': image_data,
                    'description': f"{img.format} image, {img.width}x{img.height} pixels, {img.mode} mode"
                }
                
        except Exception as e:
            raise FileProcessorError(f"Image processing failed: {str(e)}")
    
    def process_docx(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text and metadata from Word documents
        """
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            full_text = ""
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
                    full_text += para.text + "\n"
            
            # Extract text from tables
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            # Get document properties
            core_props = doc.core_properties
            metadata = {
                'author': core_props.author,
                'category': core_props.category,
                'comments': core_props.comments,
                'created': str(core_props.created) if core_props.created else None,
                'keywords': core_props.keywords,
                'language': core_props.language,
                'last_modified_by': core_props.last_modified_by,
                'modified': str(core_props.modified) if core_props.modified else None,
                'subject': core_props.subject,
                'title': core_props.title,
                'version': core_props.version
            }
            
            return {
                'type': 'docx',
                'text': full_text.strip(),
                'paragraphs': paragraphs,
                'paragraph_count': len(paragraphs),
                'tables': tables_data,
                'table_count': len(tables_data),
                'metadata': metadata,
                'text_length': len(full_text)
            }
            
        except Exception as e:
            raise FileProcessorError(f"DOCX processing failed: {str(e)}")
    
    def process_xlsx(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract data from Excel spreadsheets
        """
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            sheets_data = {}
            all_text = ""
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Get sheet data
                sheet_data = []
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    sheet_data.append(row_data)
                    all_text += " ".join(row_data) + "\n"
                
                sheets_data[sheet_name] = {
                    'data': sheet_data,
                    'rows': len(sheet_data),
                    'columns': len(sheet_data[0]) if sheet_data else 0
                }
            
            return {
                'type': 'xlsx',
                'text': all_text.strip(),
                'sheets': sheets_data,
                'sheet_count': len(workbook.sheetnames),
                'sheet_names': workbook.sheetnames,
                'text_length': len(all_text)
            }
            
        except Exception as e:
            raise FileProcessorError(f"XLSX processing failed: {str(e)}")
    
    def process_text(self, file_path: Path) -> Dict[str, Any]:
        """
        Process plain text files
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            lines = content.split('\n')
            
            return {
                'type': 'text',
                'text': content,
                'line_count': len(lines),
                'text_length': len(content),
                'encoding': 'utf-8'
            }
            
        except Exception as e:
            raise FileProcessorError(f"Text processing failed: {str(e)}")
    
    def process_csv(self, file_path: Path) -> Dict[str, Any]:
        """
        Process CSV files
        """
        try:
            import csv
            
            rows = []
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                csv_reader = csv.reader(file)
                for row in csv_reader:
                    rows.append(row)
            
            # Generate text representation
            text_content = ""
            for row in rows:
                text_content += ",".join(row) + "\n"
            
            return {
                'type': 'csv',
                'text': text_content.strip(),
                'rows': rows,
                'row_count': len(rows),
                'column_count': len(rows[0]) if rows else 0,
                'headers': rows[0] if rows else [],
                'text_length': len(text_content)
            }
            
        except Exception as e:
            raise FileProcessorError(f"CSV processing failed: {str(e)}")
    
    def process_json(self, file_path: Path) -> Dict[str, Any]:
        """
        Process JSON files
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Convert to pretty-printed text
            text_content = json.dumps(data, indent=2, ensure_ascii=False)
            
            return {
                'type': 'json',
                'text': text_content,
                'data': data,
                'text_length': len(text_content),
                'structure': self._analyze_json_structure(data)
            }
            
        except Exception as e:
            raise FileProcessorError(f"JSON processing failed: {str(e)}")
    
    def _analyze_json_structure(self, data: Any) -> Dict[str, Any]:
        """
        Analyze the structure of JSON data
        """
        if isinstance(data, dict):
            return {
                'type': 'object',
                'keys': list(data.keys()),
                'key_count': len(data)
            }
        elif isinstance(data, list):
            return {
                'type': 'array',
                'length': len(data),
                'item_types': list(set(type(item).__name__ for item in data))
            }
        else:
            return {
                'type': type(data).__name__,
                'value': str(data)
            }
    
    def get_supported_types(self) -> List[str]:
        """
        Get list of supported file types
        """
        return list(self.supported_types.keys())
    
    def get_file_summary(self, processed_data: Dict[str, Any]) -> str:
        """
        Generate a human-readable summary of processed file
        """
        if processed_data['processing_status'] != 'success':
            return f"Error processing file: {processed_data.get('error', 'Unknown error')}"
        
        content = processed_data['content']
        file_type = content['type']
        
        summaries = {
            'pdf': f"PDF document with {content.get('page_count', 0)} pages and {content.get('text_length', 0)} characters",
            'image': f"{content.get('format', 'Unknown')} image ({content.get('width', 0)}x{content.get('height', 0)}) in {content.get('mode', 'Unknown')} mode",
            'docx': f"Word document with {content.get('paragraph_count', 0)} paragraphs and {content.get('table_count', 0)} tables",
            'xlsx': f"Excel spreadsheet with {content.get('sheet_count', 0)} sheets",
            'text': f"Text file with {content.get('line_count', 0)} lines and {content.get('text_length', 0)} characters",
            'csv': f"CSV file with {content.get('row_count', 0)} rows and {content.get('column_count', 0)} columns",
            'json': f"JSON file with {content.get('text_length', 0)} characters"
        }
        
        return summaries.get(file_type, f"Processed {file_type} file")

# Convenience function
def process_file(file_path: Path) -> Dict[str, Any]:
    """
    Convenience function to process a single file
    """
    processor = FileProcessor()
    return processor.process_file(file_path)